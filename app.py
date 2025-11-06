# =============================
# ATFM Iraq – GCANS-IRAQ ADP & Visualizer
# =============================
# Features per request:
# - Session auth (any username; password: atfmiraqmm); login UI disappears after success
# - Pulls latest Google Doc (export .docx)
# - Extract & display: Airspace Information, Airport sections, NOTAMs, Weather (ATFM met forecast),
#   Predicted hourly demand (charts), ATFM measures & effects (static), CDM info (static)
# - FIR map with route animation (Tasmi→Kaban, MODIK→SIDAD)
# - Sector capacity (South/TASMI 26, RATVO 27)
# - Generate ADP (DOCX) including logo, GCANS-IRAQ provider, footer "built and supervised by MM and CU"
# =============================

import streamlit as st
import pandas as pd
import numpy as np
import requests
import re
from io import BytesIO
from datetime import datetime
from typing import Dict, List, Tuple

import plotly.express as px
import plotly.graph_objects as go
import pydeck as pdk

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# -----------------------------
# CONFIG / CONSTANTS
# -----------------------------
st.set_page_config(page_title="ATFM IRAQ – GCANS-IRAQ ADP", layout="wide")

APP_TITLE = "🇮🇶 ATFM IRAQ – GCANS-IRAQ ADP & Visualizer"
GOOGLE_DOC_EXPORT = "https://docs.google.com/document/d/1PUtfstGvw8PhKWbnOOvlBjCa7wJJX-nM/export?format=docx"

PASSWORD = "atfmiraqmm"
LOGO_URL = "https://sl.bing.net/kP2aXyPKi2C"

SECTOR_CAP = {"South/TASMI": 26, "RATVO": 27}

ROUTES = {
    "Tasmi→Kaban": {
        # approximate demo coords; replace with official if desired
        "points": [(31.3, 47.3), (33.6, 44.4), (35.55, 45.3), (36.1, 44.6)]
    },
    "MODIK→SIDAD": {
        "points": [(36.5, 42.8), (35.6, 43.7), (34.2, 44.3), (33.1, 44.5)]
    },
}

AIRPORTS = {
    "ORBI (Baghdad)": (33.2625, 44.2346),
    "ORNI (Najaf)": (31.989, 44.404),
    "ORER (Erbil)": (36.2376, 43.9632),
    "ORBS (Basra)": (30.549, 47.662),
}

# -----------------------------
# HEADER WITH LOGO
# -----------------------------
left, right = st.columns([4, 1])
with left:
    st.title(APP_TITLE)
with right:
    try:
        st.image(LOGO_URL, caption="GCANS-IRAQ", use_container_width=True)
    except Exception:
        pass

st.markdown(":blue[Service Provider: **GCANS-IRAQ**]")

# -----------------------------
# SESSION AUTH
# -----------------------------
if "auth" not in st.session_state:
    st.session_state.auth = False

if not st.session_state.auth:
    with st.form("login_form", clear_on_submit=True):
        c1, c2 = st.columns([1, 1])
        with c1:
            username = st.text_input("Username (any):")
        with c2:
            password = st.text_input("Password:", type="password")
        ok = st.form_submit_button("Login")

    if ok:
        if password == PASSWORD:
            st.session_state.auth = True
            st.experimental_rerun()
        else:
            st.error("Wrong password.")
            st.stop()
else:
    # After login, no fields shown.
    st.success("Authenticated ✅")

# -----------------------------
# FETCH LATEST WORD FILE
# -----------------------------
with st.spinner("Fetching the latest ATFM Word file…"):
    try:
        r = requests.get(GOOGLE_DOC_EXPORT, timeout=30)
        r.raise_for_status()
        DOC_BYTES = BytesIO(r.content)
        st.caption("✅ Latest document downloaded.")
    except Exception as e:
        st.error(f"❌ Failed to download the Google Doc: {e}")
        st.stop()

# -----------------------------
# PARSE DOCX → TEXT
# -----------------------------
def read_docx(buff: BytesIO) -> str:
    d = Document(buff)
    return "\n".join(p.text for p in d.paragraphs)

RAW_TEXT = read_docx(DOC_BYTES)

# -----------------------------
# GENERIC SECTION EXTRACTOR
# -----------------------------
def extract_section(text: str, keywords: List[str], stop_keywords: List[str] = None) -> str:
    """
    Extract section text starting from first line that has one of keywords
    until next line that looks like a new section (keyword2 or ALLCAPS heading)
    """
    lines = text.splitlines()
    idx_start = -1
    for i, ln in enumerate(lines):
        if any(k.lower() in ln.lower() for k in keywords):
            idx_start = i
            break
    if idx_start == -1:
        return ""

    stop_keywords = stop_keywords or []
    # heuristic for end: next heading-like line or a stop keyword
    idx_end = len(lines)
    for j in range(idx_start + 1, len(lines)):
        l = lines[j].strip()
        if any(k.lower() in l.lower() for k in stop_keywords):
            idx_end = j
            break
        # heading heuristic: very short all-caps or looks like a title
        if (len(l) < 50 and l.isupper()) and j > idx_start + 1:
            idx_end = j
            break

    return "\n".join(lines[idx_start:idx_end]).strip()

# -----------------------------
# DATA EXTRACTORS
# -----------------------------
def extract_airspace_info(text: str) -> str:
    return extract_section(
        text,
        keywords=["Airspace Information", "Airspace Status", "FIR Information"],
        stop_keywords=["Airport", "NOTAM", "Meteorological", "Weather", "Overflight", "Demand", "Capacity"]
    )

def extract_airport_traffic(text: str) -> pd.DataFrame:
    """
    Flexible arrivals/departures matcher:
    - ORBI Arrivals: 45 Departures: 50
    - ORNI A=20 D=22
    - ORER Arr 15 Dep 17
    """
    patt = re.compile(
        r"\b(ORBI|ORNI|ORER|ORBS)\b[^\n]*?"
        r"(Arr(?:ivals)?|A)\s*[:=]?\s*(\d+)?[^\n]*?"
        r"(Dep(?:artures)?|D)\s*[:=]?\s*(\d+)?",
        re.IGNORECASE
    )
    rows = []
    for m in patt.finditer(text):
        icao = m.group(1).upper()
        arr_raw = m.group(3)
        dep_raw = m.group(5)
        arr = int(arr_raw) if arr_raw and arr_raw.isdigit() else 0
        dep = int(dep_raw) if dep_raw and dep_raw.isdigit() else 0
        rows.append({"Airport": icao, "Arrivals": arr, "Departures": dep})

    if not rows:
        return pd.DataFrame(columns=["Airport", "Arrivals", "Departures"])
    return pd.DataFrame(rows).groupby("Airport", as_index=False).sum()

def extract_overflights(text: str) -> pd.DataFrame:
    """
    Lines like: 0000–0100 73   or 01:00-02:00 64
    """
    patt = re.compile(r"(\d{2}[:]?00)\s*[–\-]\s*(\d{2}[:]?00)\s+(\d+)")
    rows = []
    for m in patt.finditer(text):
        s = m.group(1).replace(":", "")
        e = m.group(2).replace(":", "")
        v = int(m.group(3))
        rows.append({"Period (UTC)": f"{s[:2]}00–{e[:2]}00", "Overflights": v})
    return pd.DataFrame(rows)

def extract_route_weather(text: str) -> Dict[str, List[Tuple[str, str]]]:
    routes = {"Tasmi→Kaban": [], "MODIK→SIDAD": []}
    lines = text.splitlines()
    current = None
    for ln in lines:
        l = ln.strip()
        if re.search(r"tasmi.*kaban", l, re.IGNORECASE):
            current = "Tasmi→Kaban"
            continue
        if re.search(r"modik.*sidad", l, re.IGNORECASE):
            current = "MODIK→SIDAD"
            continue
        if current:
            m = re.match(r"(\d{2}\s*-\s*\d{2}Z)\s*:\s*(.+)", l)
            if m:
                routes[current].append((m.group(1).replace(" ", ""), m.group(2)))

    # fallbacks if not present
    if not routes["Tasmi→Kaban"]:
        routes["Tasmi→Kaban"] = [
            ("00-06Z", "Light CAT FL200–280, W 20–25 kt"),
            ("06-12Z", "Isolated CB near route; icing above FL240"),
            ("12-18Z", "Nil SIGWX"),
            ("18-24Z", "Moderate headwind 25–30 kt"),
        ]
    if not routes["MODIK→SIDAD"]:
        routes["MODIK→SIDAD"] = [
            ("00-06Z", "Mountain wave NW, light chop"),
            ("06-12Z", "Good VIS, nil SIGWX"),
            ("12-18Z", "Convective build-ups SE of route"),
            ("18-24Z", "Crosswind shear FL180–220"),
        ]
    return routes

def extract_notams(text: str) -> List[dict]:
    lines = [l.strip() for l in text.splitlines() if l.strip()]
    cands = [l for l in lines if re.search(r"NOTAM|Q\)|RWY|RUNWAY|TWY|U/S|CLOSED|CRANE|WIP|ILS|VOR|NDB", l, re.IGNORECASE)]
    notams = []
    for l in cands:
        cat = "General"
        if re.search(r"RWY|RUNWAY", l, re.IGNORECASE): cat = "Runway"
        if re.search(r"TWY|TAXI", l, re.IGNORECASE):   cat = "Taxiway"
        if re.search(r"ILS|VOR|NDB", l, re.IGNORECASE): cat = "NavAid"
        if re.search(r"CLOSED|U/S|UNSERVICEABLE", l, re.IGNORECASE): cat = "Closure/Unserviceable"
        if re.search(r"CRANE|WIP|WORK IN PROGRESS", l, re.IGNORECASE): cat = "Obstacles/Work"
        icao = "-"
        m = re.search(r"\bOR[A-Z]{2}\b", l)
        if m: icao = m.group(0)
        notams.append({"Airport": icao, "Category": cat, "Text": l})
    return notams

# -----------------------------
# RUN EXTRACTION
# -----------------------------
airspace_info = extract_airspace_info(RAW_TEXT)
airport_df     = extract_airport_traffic(RAW_TEXT)
overflights_df = extract_overflights(RAW_TEXT)
route_weather  = extract_route_weather(RAW_TEXT)
notams         = extract_notams(RAW_TEXT)

# -----------------------------
# ATFM MEASURES & EFFECTS (STATIC DAILY)
# -----------------------------
ATFM_MEASURES = [
    ("Miles-in-Trail (MIT) / Minutes-in-Trail (MINIT)", "Reduces sector entry rate and stream compression; smooths demand but increases delay and elongates trajectories."),
    ("Level Capping / RALT", "Avoids conflicted levels or congested layers; may increase fuel burn but preserves sector capacity and separation."),
    ("Rerouting (e.g., via TASMI→KABAN only)", "Shifts flows away from hot spots; may shift demand to adjacent sectors and add track miles."),
    ("GDP (Ground Delay Program)", "Holds departures at origin to meter arrivals/overflights; delays on ground instead of airborne holding."),
    ("GS (Ground Stop)", "Temporary stoppage of departures to affected FIR/airport; used for severe capacity drops or emergencies."),
    ("Rate Regulation (MAX xx/hr)", "Caps acceptance rate to sector capacity; predictable but can propagate delay upstream."),
]

CDM_INFO = [
    "Shared data between ACC/ATFMU, airlines, airports, MET, and adjacent FIRs.",
    "Common situational picture (traffic, capacity, MET, NOTAM, constraints).",
    "Collaborative assessment → agreed regulation (scope, duration, thresholds).",
    "Continuous review: adjust measure when demand/capacity changes.",
]

# -----------------------------
# LAYOUT SECTIONS
# -----------------------------
st.markdown("## 🛰️ Airspace Information")
if airspace_info.strip():
    st.write(airspace_info)
else:
    st.info("No explicit *Airspace Information* section found. Add a heading like 'Airspace Information' in the document to auto-extract.")

st.markdown("## 🛬 Airport Information (one by one)")
if airport_df.empty:
    st.info("No airport figures found. Use patterns like `ORBI Arrivals: 45 Departures: 50` or `ORNI A=20 D=22` in your document.")
else:
    for ap in ["ORBI", "ORNI", "ORER", "ORBS"]:
        part = airport_df[airport_df["Airport"] == ap]
        with st.expander(f"{ap} details", expanded=True if ap == "ORBI" else False):
            if part.empty:
                st.write("No data.")
            else:
                st.dataframe(part, use_container_width=True)
                arr = int(part["Arrivals"].sum())
                dep = int(part["Departures"].sum())
                st.metric("Arrivals", arr)
                st.metric("Departures", dep)
                # little bar
                md = part.melt("Airport", var_name="Type", value_name="Flights")
                fig_ap = px.bar(md, x="Type", y="Flights", title=f"{ap} Traffic")
                st.plotly_chart(fig_ap, use_container_width=True)

st.markdown("## ⏱️ Predicted Demand (Hourly Overflights)")
if overflights_df.empty:
    st.info("No hourly overflight lines found (e.g., `0000–0100 73`).")
else:
    st.dataframe(overflights_df, use_container_width=True)
    fig_line = px.area(overflights_df, x="Period (UTC)", y="Overflights", title="Predicted Hourly Demand (Overflights)", markers=True)
    st.plotly_chart(fig_line, use_container_width=True)
    peak = int(overflights_df["Overflights"].max())
    off_peak = int(overflights_df["Overflights"].median())
    st.caption(f"Peak hour demand: **{peak}** | Median hour: **{off_peak}**")

st.markdown("## 📈 Sector Capacity & Utilization")
est_peak = int(overflights_df["Overflights"].max()) if not overflights_df.empty else 20
cap_rows = []
for sec, cap in SECTOR_CAP.items():
    util = round((est_peak / cap) * 100, 1)
    cap_rows.append({"Sector": sec, "Capacity (acft/hr)": cap, "Estimated Peak Demand": est_peak, "Utilization %": util})
cap_df = pd.DataFrame(cap_rows)
c1, c2 = st.columns([2, 1])
with c1:
    st.dataframe(cap_df, use_container_width=True)
with c2:
    gauge = go.Figure()
    for _, r in cap_df.iterrows():
        gauge.add_trace(go.Indicator(mode="gauge+number",
                                     value=r["Utilization %"],
                                     title={"text": r["Sector"]},
                                     gauge={"axis": {"range": [0, 150]}}))
    gauge.update_layout(height=380)
    st.plotly_chart(gauge, use_container_width=True)

st.markdown("## 🌤️ ATFM Meteorological Forecast (Routes)")
route_tabs = st.tabs(["Tasmi→Kaban", "MODIK→SIDAD"])
def weather_viz(route_key: str):
    data = route_weather.get(route_key, [])
    wdf = pd.DataFrame(data, columns=["Window", "Summary"])
    # Derive a simple severity for visualization
    def sev(s: str) -> int:
        s = s.lower()
        score = 1
        if "cb" in s or "convect" in s: score += 3
        if "turb" in s or "chop" in s or "cat" in s: score += 2
        if "icing" in s: score += 2
        if "headwind" in s or "crosswind" in s or "shear" in s: score += 1
        return min(5, score)
    if not wdf.empty:
        wdf["Severity"] = wdf["Summary"].apply(sev)
        st.dataframe(wdf, use_container_width=True)
        fig = px.bar(wdf, x="Window", y="Severity", text="Summary",
                     animation_frame="Window", range_y=[0, 6],
                     title=f"Weather Severity – {route_key}")
        fig.update_traces(textposition="outside")
        st.plotly_chart(fig, use_container_width=True)
        st.caption("Explanation: Severity index is a heuristic to visualize changing risk across the day (1 low → 5 high).")
    else:
        st.info("No route weather windows found; using daily defaults.")

with route_tabs[0]:
    weather_viz("Tasmi→Kaban")
with route_tabs[1]:
    weather_viz("MODIK→SIDAD")

# -----------------------------
# NOTAMs
# -----------------------------
st.markdown("## 📜 NOTAMs – Interpretation")
if not notams:
    st.info("No NOTAM-like lines detected.")
else:
    notam_df = pd.DataFrame(notams)
    st.dataframe(notam_df, use_container_width=True)
    sum_fig = px.bar(notam_df.groupby("Category").size().reset_index(name="Count"),
                     x="Category", y="Count", title="NOTAMs by Category")
    st.plotly_chart(sum_fig, use_container_width=True)

# -----------------------------
# FIR MAP + ROUTE ANIMATION
# -----------------------------
st.markdown("## 🗺️ Iraqi FIR Map + Route Animation")
center = [33.3, 44.4]
layers = []
# Airports
airport_points = [{"name": k, "lat": v[0], "lon": v[1]} for k, v in AIRPORTS.items()]
layers.append(pdk.Layer("ScatterplotLayer", data=airport_points,
                        get_position=["lon", "lat"], get_radius=15000, get_color=[255, 0, 0], pickable=True))
# Route lines
route_lines = []
for name, info in ROUTES.items():
    pts = info["points"]
    for i in range(len(pts) - 1):
        route_lines.append({"name": name, "from_lon": pts[i][1], "from_lat": pts[i][0],
                            "to_lon": pts[i+1][1], "to_lat": pts[i+1][0]})
layers.append(pdk.Layer("LineLayer", data=route_lines,
                        get_source_position=["from_lon", "from_lat"],
                        get_target_position=["to_lon", "to_lat"],
                        get_width=3, get_color=[0, 100, 255], pickable=True))

st.caption("Move the slider to animate a probe aircraft along the selected route.")
rsel = st.selectbox("Route", list(ROUTES.keys()), index=0)
prog = st.slider("Progress", 0, 100, 0)

def interpolate(points, t):
    if t <= 0: return points[0]
    if t >= 1: return points[-1]
    seg = int(t * (len(points) - 1))
    seg = min(seg, len(points) - 2)
    local_t = t * (len(points) - 1) - seg
    lat1, lon1 = points[seg]; lat2, lon2 = points[seg+1]
    return (lat1 + (lat2-lat1)*local_t, lon1 + (lon2-lon1)*local_t)

ilat, ilon = interpolate(ROUTES[rsel]["points"], prog/100.0)
layers.append(pdk.Layer("ScatterplotLayer", data=[{"lat": ilat, "lon": ilon, "name": "Probe"}],
                        get_position=["lon", "lat"], get_radius=20000, get_color=[0, 255, 0], pickable=True))

st.pydeck_chart(pdk.Deck(map_style="light",
                         initial_view_state=pdk.ViewState(latitude=center[0], longitude=center[1], zoom=5),
                         layers=layers, tooltip={"text": "{name}"}))

# -----------------------------
# ATFM MEASURES (STATIC DAILY)
# -----------------------------
st.markdown("## 🧰 ATFM Measures (Applied/Available) – Effects")
for name, effect in ATFM_MEASURES:
    st.markdown(f"- **{name}** — {effect}")

# -----------------------------
# CDM INFO (ADDITIONAL)
# -----------------------------
st.markdown("## 🤝 CDM (Collaborative Decision Making) – Daily Guidance")
for item in CDM_INFO:
    st.markdown(f"- {item}")

# -----------------------------
# ADP GENERATOR (DOCX)
# -----------------------------
st.markdown("## 🧾 Generate ADP (DOCX) – GCANS-IRAQ")

adp_date = st.date_input("ADP Date (UTC)", value=datetime.utcnow().date())

def build_adp_docx() -> BytesIO:
    doc = Document()

    # Try to add logo
    try:
        lr = requests.get(LOGO_URL, timeout=10)
        if lr.ok:
            doc.add_picture(BytesIO(lr.content), width=Inches(1.2))
    except Exception:
        pass

    title = doc.add_paragraph(f"ATFM Daily Plan – {adp_date.isoformat()}")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].bold = True
    title.runs[0].font.size = Pt(16)

    prov = doc.add_paragraph("Service Provider: GCANS-IRAQ")
    prov.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(f"Generated: {datetime.utcnow().strftime('%Y-%m-%d %H:%MZ')}")

    # Airspace info
    doc.add_paragraph().add_run("Airspace Information").bold = True
    if airspace_info.strip():
        doc.add_paragraph(airspace_info)
    else:
        doc.add_paragraph("N/A")

    # Airport tables
    doc.add_paragraph().add_run("Airports").bold = True
    if airport_df.empty:
        doc.add_paragraph("No airport figures detected.")
    else:
        t = doc.add_table(rows=1, cols=3)
        hdr = t.rows[0].cells
        hdr[0].text = "Airport"
        hdr[1].text = "Arrivals"
        hdr[2].text = "Departures"
        for _, r in airport_df.iterrows():
            row = t.add_row().cells
            row[0].text = r["Airport"]
            row[1].text = str(int(r["Arrivals"]))
            row[2].text = str(int(r["Departures"]))

    # Predicted demand
    doc.add_paragraph().add_run("Predicted Demand (Hourly Overflights)").bold = True
    if overflights_df.empty:
        doc.add_paragraph("No hourly overflight lines found.")
    else:
        t = doc.add_table(rows=1, cols=2)
        hdr = t.rows[0].cells
        hdr[0].text = "Period (UTC)"
        hdr[1].text = "Overflights"
        for _, r in overflights_df.iterrows():
            row = t.add_row().cells
            row[0].text = r["Period (UTC)"]
            row[1].text = str(int(r["Overflights"]))

    # Capacity
    doc.add_paragraph().add_run("Sector Capacity & Utilization").bold = True
    t = doc.add_table(rows=1, cols=4)
    hdr = t.rows[0].cells
    hdr[0].text = "Sector"
    hdr[1].text = "Capacity (acft/hr)"
    hdr[2].text = "Estimated Peak Demand"
    hdr[3].text = "Utilization"
    for _, r in cap_df.iterrows():
        row = t.add_row().cells
        row[0].text = r["Sector"]
        row[1].text = str(int(r["Capacity (acft/hr)"]))
        row[2].text = str(int(r["Estimated Peak Demand"]))
        row[3].text = f"{float(r['Utilization %']):.0f}%"

    # NOTAMs
    doc.add_paragraph().add_run("NOTAMs – Summary").bold = True
    if not notams:
        doc.add_paragraph("No NOTAM-like items found.")
    else:
        for n in notams:
            doc.add_paragraph(f"- [{n['Category']}] {n['Airport']}: {n['Text']}")

    # Weather routes
    doc.add_paragraph().add_run("ATFM Meteorological Forecast").bold = True
    for rname, items in route_weather.items():
        doc.add_paragraph(rname)
        for win, summ in items:
            doc.add_paragraph(f"  • {win}: {summ}")

    # ATFM measures (static)
    doc.add_paragraph().add_run("ATFM Measures & Effects (Daily Reference)").bold = True
    for name, eff in ATFM_MEASURES:
        doc.add_paragraph(f"- {name}: {eff}")

    # CDM info
    doc.add_paragraph().add_run("CDM (Collaborative Decision Making) – Notes").bold = True
    for it in CDM_INFO:
        doc.add_paragraph(f"- {it}")

    # Footer
    foot = doc.add_paragraph("This app built and supervised by MM and CU")
    foot.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    out = BytesIO()
    doc.save(out)
    out.seek(0)
    return out

if st.button("Generate ADP (DOCX)"):
    file = build_adp_docx()
    st.download_button(
        "⬇️ Download ADP (DOCX)",
        data=file,
        file_name=f"ATFM_ADP_{adp_date.isoformat()}.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

# -----------------------------
# FOOTER
# -----------------------------
st.divider()
st.caption("© GCANS-IRAQ — This app built and supervised by **MM** and **CU**")
